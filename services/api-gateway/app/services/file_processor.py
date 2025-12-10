"""
File processing service for OCR, text extraction, and document parsing
"""

import mimetypes
from io import BytesIO
from typing import Optional, Tuple

from app.core.logging import get_logger

logger = get_logger(__name__)


class FileProcessor:
    """Service for processing uploaded files and extracting text content"""

    def __init__(self):
        """Initialize file processor with optional OCR support"""
        self.ocr_available = False
        try:
            import pytesseract  # noqa: F401

            self.ocr_available = True
            logger.info("OCR support enabled (pytesseract available)")
        except ImportError:
            logger.warning("OCR support disabled (pytesseract not installed)")

    async def extract_text(
        self, file_content: bytes, filename: str, mime_type: Optional[str] = None
    ) -> Tuple[str, dict]:
        """
        Extract text from a file.

        Args:
            file_content: File content as bytes
            filename: Original filename
            mime_type: MIME type of file

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        if not mime_type:
            mime_type, _ = mimetypes.guess_type(filename)

        metadata = {
            "filename": filename,
            "mime_type": mime_type,
            "size_bytes": len(file_content),
        }

        try:
            if mime_type == "application/pdf":
                text, pdf_metadata = await self._extract_from_pdf(file_content)
                metadata.update(pdf_metadata)
                return text, metadata

            elif mime_type in [
                "image/png",
                "image/jpeg",
                "image/jpg",
                "image/tiff",
                "image/bmp",
            ]:
                text, img_metadata = await self._extract_from_image(file_content)
                metadata.update(img_metadata)
                return text, metadata

            elif mime_type == "text/plain":
                text = file_content.decode("utf-8", errors="ignore")
                return text, metadata

            elif mime_type == "text/markdown":
                text = file_content.decode("utf-8", errors="ignore")
                metadata["format"] = "markdown"
                return text, metadata

            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                text, doc_metadata = await self._extract_from_docx(file_content)
                metadata.update(doc_metadata)
                return text, metadata

            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                return "", metadata

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}", exc_info=True)
            metadata["error"] = str(e)
            return "", metadata

    async def _extract_from_pdf(self, file_content: bytes) -> Tuple[str, dict]:
        """Extract text from PDF file"""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return "", {"error": "PDF support not available"}

        try:
            pdf_file = BytesIO(file_content)
            reader = PdfReader(pdf_file)

            metadata = {
                "page_count": len(reader.pages),
                "format": "pdf",
            }

            # Extract PDF metadata
            if reader.metadata:
                metadata["title"] = reader.metadata.get("/Title", "")
                metadata["author"] = reader.metadata.get("/Author", "")
                metadata["subject"] = reader.metadata.get("/Subject", "")

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {e}")

            text = "\n\n".join(text_parts)
            return text, metadata

        except Exception as e:
            logger.error(f"Error processing PDF: {e}", exc_info=True)
            return "", {"error": str(e)}

    async def _extract_from_image(self, file_content: bytes) -> Tuple[str, dict]:
        """Extract text from image using OCR"""
        if not self.ocr_available:
            return "", {"error": "OCR not available (pytesseract not installed)"}

        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            logger.error("Image processing libraries not installed. " "Install with: pip install Pillow pytesseract")
            return "", {"error": "Image processing not available"}

        try:
            image = Image.open(BytesIO(file_content))

            metadata = {
                "format": "image",
                "width": image.width,
                "height": image.height,
                "mode": image.mode,
            }

            # Perform OCR
            text = pytesseract.image_to_string(image)
            metadata["ocr_performed"] = True

            return text, metadata

        except Exception as e:
            logger.error(f"Error processing image: {e}", exc_info=True)
            return "", {"error": str(e)}

    async def _extract_from_docx(self, file_content: bytes) -> Tuple[str, dict]:
        """Extract text from DOCX file"""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return "", {"error": "DOCX support not available"}

        try:
            doc = Document(BytesIO(file_content))

            metadata = {
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
            }

            # Extract document properties
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(row_text))

            text = "\n\n".join(text_parts)
            return text, metadata

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}", exc_info=True)
            return "", {"error": str(e)}

    def get_supported_types(self) -> list:
        """Get list of supported MIME types"""
        supported = [
            "application/pdf",
            "text/plain",
            "text/markdown",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

        if self.ocr_available:
            supported.extend(
                [
                    "image/png",
                    "image/jpeg",
                    "image/jpg",
                    "image/tiff",
                    "image/bmp",
                ]
            )

        return supported

    async def validate_file(
        self, file_content: bytes, filename: str, max_size_mb: int = 10
    ) -> Tuple[bool, Optional[str]]:
        """
        Validate file before processing.

        Args:
            file_content: File content as bytes
            filename: Original filename
            max_size_mb: Maximum file size in MB

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        size_mb = len(file_content) / (1024 * 1024)
        if size_mb > max_size_mb:
            return False, f"File too large ({size_mb:.1f}MB). Maximum: {max_size_mb}MB"

        # Check file type
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type not in self.get_supported_types():
            return (
                False,
                f"Unsupported file type: {mime_type}. " f"Supported types: {', '.join(self.get_supported_types())}",
            )

        # Check if file is empty
        if len(file_content) == 0:
            return False, "File is empty"

        return True, None


# Singleton instance
_file_processor = None


def get_file_processor() -> FileProcessor:
    """Get file processor instance"""
    global _file_processor
    if _file_processor is None:
        _file_processor = FileProcessor()
    return _file_processor
